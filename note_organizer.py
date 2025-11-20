import tkinter as tk
from tkinter import filedialog, scrolledtext, messagebox, Menu, Toplevel, ttk
import requests
import json
import os
import threading
import base64
import io
import re
from datetime import datetime

# Check for optional libraries
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pypdf
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from PIL import Image, ImageTk
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

# Themes Configuration
THEMES = {
    "Day Mode": {
        "bg": "#f0f2f5", "fg": "#1f2937", "frame_bg": "#ffffff",
        "text_bg": "#f9fafb", "text_fg": "#000000",
        "btn_primary": "#4f46e5", "btn_secondary": "#ef4444", "btn_success": "#10b981", "btn_fg": "#ffffff"
    },
    "Night Mode": {
        "bg": "#111827", "fg": "#e5e7eb", "frame_bg": "#1f2937",
        "text_bg": "#374151", "text_fg": "#ffffff",
        "btn_primary": "#6366f1", "btn_secondary": "#ef4444", "btn_success": "#10b981", "btn_fg": "#ffffff"
    }
}

# Limits
MAX_RAW_CHARS = 200_000  # warn/truncate above this
MAX_IMAGE_PIXELS = 1200  # max width/height when downscaling images

class NoteOrganizerApp:
    def __init__(self, master):
        self.master = master
        master.title("Note Organizer")
        master.geometry("950x850")
        
        # Default to Day Mode, but load_config will override this if JSON says otherwise
        self.current_theme_name = "Day Mode"
        self.colors = THEMES[self.current_theme_name]

        self.api_key = ""
        self.api_url = ""
        
        # --- LOAD CONFIGURATION ---
        self.load_config()

        self.raw_text_content = ""
        self.image_payloads = []
        self.loaded_files_list = []

        self.create_menu_bar()

        master.grid_columnconfigure(0, weight=1)
        master.grid_rowconfigure(2, weight=1)

        # Header
        self.header_frame = tk.Frame(master)
        self.header_frame.grid(row=0, column=0, sticky="ew", padx=10, pady=5)
        
        self.title_label = tk.Label(self.header_frame, text="Note Organizer", font=("Helvetica", 18, "bold"))
        self.title_label.pack(side=tk.LEFT)
        
        self.subtitle_label = tk.Label(self.header_frame, text="(Categorization + Tables)", font=("Helvetica", 10))
        self.subtitle_label.pack(side=tk.LEFT, padx=10)

        # Main Container
        self.main_frame = tk.Frame(master, padx=10, pady=10)
        self.main_frame.grid(row=2, column=0, sticky="nsew")
        self.main_frame.grid_columnconfigure(0, weight=1)
        self.main_frame.grid_columnconfigure(1, weight=1)
        self.main_frame.grid_rowconfigure(0, weight=1)

        # --- LEFT COLUMN: INPUT ---
        self.input_frame = tk.Frame(self.main_frame, padx=10, pady=10, relief=tk.GROOVE, bd=1)
        self.input_frame.grid(row=0, column=0, sticky="nsew", padx=5, pady=5)
        self.input_frame.grid_columnconfigure(0, weight=1)
        self.input_frame.grid_rowconfigure(2, weight=1)

        self.lbl_step1 = tk.Label(self.input_frame, text="1. Source Materials", font=("Helvetica", 12, "bold"))
        self.lbl_step1.grid(row=0, column=0, sticky="w")
        
        self.btn_frame = tk.Frame(self.input_frame)
        self.btn_frame.grid(row=1, column=0, sticky="ew", pady=5)
        
        self.btn_upload = tk.Button(self.btn_frame, text="Add Files", command=self.upload_files, relief=tk.FLAT, padx=10)
        self.btn_upload.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 2))
        
        self.btn_clear = tk.Button(self.btn_frame, text="Reset", command=self.clear_inputs, relief=tk.FLAT)
        self.btn_clear.pack(side=tk.LEFT, padx=(2, 0))

        self.lbl_context = tk.Label(self.input_frame, text="Loaded Files:", anchor="w")
        self.lbl_context.grid(row=2, column=0, sticky="w", pady=(10, 0))
        
        self.file_list_display = scrolledtext.ScrolledText(self.input_frame, wrap=tk.WORD, height=8, font=("Consolas", 9))
        self.file_list_display.grid(row=3, column=0, sticky="nsew")
        self.file_list_display.insert(tk.END, "No files loaded.")
        self.file_list_display.config(state=tk.DISABLED)

        self.lbl_instruct = tk.Label(self.input_frame, text="Instructions:", anchor="w")
        self.lbl_instruct.grid(row=4, column=0, sticky="w", pady=(10, 0))
        
        self.user_prompt_text = tk.Text(self.input_frame, wrap=tk.WORD, height=4, font=("Helvetica", 10))
        self.user_prompt_text.grid(row=5, column=0, sticky="ew", pady=(0, 10))

        self.process_button = tk.Button(self.input_frame, text="Categorize & Compile", command=self.start_processing_thread, font=("Helvetica", 11, "bold"), relief=tk.FLAT, pady=8)
        self.process_button.grid(row=6, column=0, sticky="ew")

        # --- RIGHT COLUMN: OUTPUT ---
        self.output_frame = tk.Frame(self.main_frame, padx=10, pady=10, relief=tk.GROOVE, bd=1)
        self.output_frame.grid(row=0, column=1, sticky="nsew", padx=5, pady=5)
        self.output_frame.grid_columnconfigure(0, weight=1)
        self.output_frame.grid_rowconfigure(1, weight=1)

        self.lbl_step2 = tk.Label(self.output_frame, text="2. Categorized Output", font=("Helvetica", 12, "bold"))
        self.lbl_step2.grid(row=0, column=0, sticky="w")

        self.compiled_output_text = scrolledtext.ScrolledText(self.output_frame, wrap=tk.WORD, font=("Helvetica", 10))
        self.compiled_output_text.grid(row=1, column=0, sticky="nsew")
        self.compiled_output_text.insert(tk.END, "Notes will appear here.")
        self.compiled_output_text.config(state=tk.DISABLED)

        self.action_frame = tk.Frame(self.output_frame)
        self.action_frame.grid(row=2, column=0, sticky="ew", pady=5)
        
        self.btn_copy = tk.Button(self.action_frame, text="Copy", command=self.copy_output)
        self.btn_copy.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=2)
        
        self.docx_btn = tk.Button(self.action_frame, text="Export DOCX", command=self.export_to_docx)
        self.docx_btn.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=2)

        # Status Bar
        self.status_label = tk.Label(master, text="Ready.", bd=1, relief=tk.SUNKEN, anchor=tk.W, pady=5, padx=5)
        self.status_label.grid(row=3, column=0, sticky="ew")

        # Progressbar
        self.progress = ttk.Progressbar(master, mode='indeterminate')
        self.progress.grid(row=4, column=0, sticky="ew", padx=10, pady=(0,10))
        self.progress.grid_remove()

        # Final Init Steps
        self.apply_theme()
        self.check_dependencies()

    def load_config(self):
        """Loads config.json from the script's directory to ensure reliability."""
        cfg = {}
        try:
            # 1. Determine the absolute path to the folder where this script lives
            script_dir = os.path.dirname(os.path.abspath(__file__))
            config_path = os.path.join(script_dir, "config.json")

            # 2. Load config from that explicit path
            if os.path.exists(config_path):
                with open(config_path, "r", encoding="utf-8") as f:
                    cfg = json.load(f)
            else:
                print(f"Warning: config.json not found at {config_path}")

            # API key priority: config file, then environment variables
            self.api_key = cfg.get("GEMINI_API_KEY") or cfg.get("API_KEY") or os.environ.get("GEMINI_API_KEY") or os.environ.get("API_KEY") or ""
            
            # Get model name, default to a stable version if missing
            model_name = cfg.get("GEMINI_MODEL", "gemini-2.0-flash")
            self.api_url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:generateContent"
            
            # Apply Theme from Config if present
            theme = cfg.get("APP_THEME")
            if theme in THEMES:
                self.current_theme_name = theme
                self.colors = THEMES[theme]
                
        except Exception as e:
            print(f"Config Loading Error: {e}")
            # Fallback if everything fails
            self.api_key = os.environ.get("GEMINI_API_KEY", "")

    def create_menu_bar(self):
        menubar = Menu(self.master)
        filemenu = Menu(menubar, tearoff=0)
        filemenu.add_command(label="Add Files...", command=self.upload_files)
        filemenu.add_command(label="Save DOCX...", command=self.export_to_docx)
        filemenu.add_separator()
        filemenu.add_command(label="Exit", command=self.master.quit)
        menubar.add_cascade(label="File", menu=filemenu)
        
        settingsmenu = Menu(menubar, tearoff=0)
        settingsmenu.add_command(label="Theme Settings", command=self.open_settings_window)
        menubar.add_cascade(label="Settings", menu=settingsmenu)
        
        self.master.config(menu=menubar)

    def open_settings_window(self):
        settings_win = Toplevel(self.master)
        settings_win.title("Settings")
        settings_win.geometry("300x220")
        settings_win.configure(bg=self.colors["bg"])
        
        tk.Label(settings_win, text="App Theme", font=("Helvetica", 12, "bold"), bg=self.colors["bg"], fg=self.colors["fg"]).pack(pady=10)
        theme_var = tk.StringVar(value=self.current_theme_name)
        
        def set_theme():
            self.current_theme_name = theme_var.get()
            self.colors = THEMES[self.current_theme_name]
            self.apply_theme()
            # Save preference to config.json if possible
            try:
                script_dir = os.path.dirname(os.path.abspath(__file__))
                config_path = os.path.join(script_dir, "config.json")
                if os.path.exists(config_path):
                    with open(config_path,"r",encoding="utf-8") as f:
                        cfg = json.load(f)
                    cfg["APP_THEME"] = self.current_theme_name
                    with open(config_path,"w",encoding="utf-8") as f:
                        json.dump(cfg, f, indent=2, ensure_ascii=False)
            except Exception as e:
                print(f"Error saving theme: {e}")
            settings_win.destroy()

        for theme_name in THEMES.keys():
            tk.Radiobutton(settings_win, text=theme_name, variable=theme_var, value=theme_name, bg=self.colors["bg"], fg=self.colors["fg"], selectcolor=self.colors["frame_bg"]).pack(anchor="w", padx=50)
        
        tk.Button(settings_win, text="Apply", command=set_theme, bg=self.colors["btn_primary"], fg="white").pack(pady=15)

    def apply_theme(self):
        c = self.colors
        self.master.config(bg=c["bg"])
        self.status_label.config(bg=c["frame_bg"], fg=c["fg"])
        self.header_frame.config(bg=c["bg"])
        self.title_label.config(bg=c["bg"], fg=c["fg"])
        self.subtitle_label.config(bg=c["bg"], fg="gray")
        self.main_frame.config(bg=c["bg"])
        self.input_frame.config(bg=c["frame_bg"])
        self.output_frame.config(bg=c["frame_bg"])
        self.btn_frame.config(bg=c["frame_bg"])
        self.action_frame.config(bg=c["frame_bg"])
        
        for lbl in [self.lbl_step1, self.lbl_step2, self.lbl_context, self.lbl_instruct]:
            lbl.config(bg=c["frame_bg"], fg=c["fg"])
            
        self.file_list_display.config(bg=c["text_bg"], fg=c["text_fg"])
        self.user_prompt_text.config(bg=c["text_bg"], fg=c["text_fg"], insertbackground=c["fg"])
        self.compiled_output_text.config(bg=c["text_bg"], fg=c["text_fg"], insertbackground=c["fg"])
        self.btn_upload.config(bg=c["btn_primary"], fg=c["btn_fg"])
        self.btn_clear.config(bg=c["btn_secondary"], fg=c["btn_fg"])
        self.process_button.config(bg=c.get("btn_success", c.get("btn_primary")), fg="white")
        self.btn_copy.config(bg=c["btn_primary"], fg=c["btn_fg"])
        self.docx_btn.config(bg=c["btn_primary"], fg=c["btn_fg"])

    def check_dependencies(self):
        if not DOCX_AVAILABLE or not PDF_AVAILABLE or not PIL_AVAILABLE:
            missing = []
            if not DOCX_AVAILABLE: missing.append("python-docx")
            if not PDF_AVAILABLE: missing.append("pypdf")
            if not PIL_AVAILABLE: missing.append("pillow")
            messagebox.showwarning("Missing Libraries", "Missing: " + ", ".join(missing))
            if not DOCX_AVAILABLE: self.docx_btn.config(state=tk.DISABLED)

    def upload_files(self):
        filepaths = filedialog.askopenfilenames(filetypes=[("All", "*.*"), ("Text", "*.txt"), ("PDF", "*.pdf"), ("Images", "*.png *.jpg *.jpeg")])
        if not filepaths: return

        for path in filepaths:
            filename = os.path.basename(path)
            ext = os.path.splitext(filename)[1].lower()
            try:
                if ext == ".txt":
                    with open(path, "r", encoding="utf-8", errors="replace") as f:
                        data = f.read()
                        self.raw_text_content += f"\n\n--- FILE: {filename} ---\n{data}"
                        self.loaded_files_list.append(filename)
                elif ext == ".pdf" and PDF_AVAILABLE:
                    reader = pypdf.PdfReader(path)
                    text_parts = []
                    for p in reader.pages:
                        try:
                            t = p.extract_text()
                        except Exception:
                            t = None
                        text_parts.append(t or "")
                    text = "\n".join(text_parts)
                    self.raw_text_content += f"\n\n--- PDF: {filename} ---\n{text}"
                    self.loaded_files_list.append(filename)
                elif ext in [".png", ".jpg", ".jpeg"] and PIL_AVAILABLE:
                    with Image.open(path) as img:
                        # downscale to reduce size
                        img.thumbnail((MAX_IMAGE_PIXELS, MAX_IMAGE_PIXELS))
                        if img.mode != "RGB":
                            img = img.convert("RGB")
                        byte_arr = io.BytesIO()
                        img.save(byte_arr, format='JPEG', quality=85)
                        encoded = base64.b64encode(byte_arr.getvalue()).decode('utf-8')
                        self.image_payloads.append({"inline_data": {"mime_type": "image/jpeg", "data": encoded}})
                        self.loaded_files_list.append(filename)
                else:
                    # unsupported file types: still add name
                    self.loaded_files_list.append(filename)
            except Exception as e:
                self.loaded_files_list.append(f"Error: {filename}")

        # if raw text is huge, warn
        if len(self.raw_text_content) > MAX_RAW_CHARS:
            if messagebox.askyesno("Large input", "Loaded text is large and may exceed model limits. Truncate to first 200k chars? (recommended) "):
                self.raw_text_content = self.raw_text_content[:MAX_RAW_CHARS]

        self.update_file_display()
        self.status_label.config(text=f"Files loaded: {len(self.loaded_files_list)}")

    def clear_inputs(self):
        self.raw_text_content = ""
        self.image_payloads = []
        self.loaded_files_list = []
        self.update_file_display()
        self.status_label.config(text="Inputs cleared.")

    def update_file_display(self):
        self.file_list_display.config(state=tk.NORMAL)
        self.file_list_display.delete("1.0", tk.END)
        if self.loaded_files_list:
            self.file_list_display.insert(tk.END, "\n".join(self.loaded_files_list))
        else:
            self.file_list_display.insert(tk.END, "No files loaded.")
        self.file_list_display.config(state=tk.DISABLED)

    def start_processing_thread(self):
        if not self.api_key:
            messagebox.showerror("Error", "API Key missing. Please check config.json or set GEMINI_API_KEY environment variable.")
            return
        
        if not self.raw_text_content and not self.image_payloads:
            messagebox.showinfo("Info", "Please upload files first.")
            return

        # disable UI elements to avoid concurrent changes
        self.process_button.config(state=tk.DISABLED, text="Processing...")
        self.btn_upload.config(state=tk.DISABLED)
        self.btn_clear.config(state=tk.DISABLED)
        self.progress.grid()
        self.progress.start(10)
        self.status_label.config(text="Organizing notes...", fg=self.colors.get("fg"))

        t = threading.Thread(target=self.process_with_gemini, daemon=True)
        t.start()

    def process_with_gemini(self):
        try:
            user_instr = self.user_prompt_text.get("1.0", tk.END).strip()
            system_instruction = (
                "Act as a professional Academic Editor. "
                "Your goal is to restructure the raw input into a Master Study Guide. "
                "1. SEGMENTATION: Organize content into distinct thematic categories. "
                "2. TABLES: You MUST use Markdown tables for any comparative data, dates, pros/cons, or formulas. "
                "3. FORMATTING: Use H2 (##) for Categories and H3 (###) for sub-topics. "
                "4. GLOSSARY: End with a glossary of key terms."
            )

            prompt_text = (
                "Restructure the following input into a categorized study guide with tables.\n"
                f"User Instructions: {user_instr}\n\n"
                "--- RAW DATA ---\n"
                f"{self.raw_text_content}\n"
                "--- END RAW DATA ---\n"
            )
            
            parts = [{"text": prompt_text}]
            parts.extend(self.image_payloads)

            payload = {
                "contents": [{"parts": parts}],
                "systemInstruction": {"parts": [{"text": system_instruction}]},
            }

            headers = {
                "Content-Type": "application/json",
            }

            # Add query param for API key
            url_with_key = f"{self.api_url}?key={self.api_key}"

            # Use json= to let `requests` set content-type properly
            resp = requests.post(url_with_key, headers=headers, json=payload, timeout=120)
            
            if resp.status_code != 200:
                raise Exception(f"API Error {resp.status_code}: {resp.text}")

            result = resp.json()

            # tolerant parsing of several plausible response shapes
            output_text = None
            try:
                output_text = result.get('candidates', [{}])[0].get('content', {}).get('parts', [{}])[0].get('text')
            except Exception:
                output_text = None

            if not output_text:
                # fallback shapes
                output_text = (result.get('choices', [{}])[0].get('message', {}).get('content') or
                               result.get('text') or json.dumps(result))

            self.master.after(0, self.finish_processing, output_text or "No content generated.")

        except Exception as e:
            self.master.after(0, self.finish_processing, f"Error: {str(e)}")
        finally:
            def restore_ui():
                self.process_button.config(state=tk.NORMAL, text="Categorize & Compile")
                self.btn_upload.config(state=tk.NORMAL)
                self.btn_clear.config(state=tk.NORMAL)
                self.progress.stop()
                self.progress.grid_remove()
            self.master.after(0, restore_ui)

    def finish_processing(self, text):
        self.compiled_output_text.config(state=tk.NORMAL)
        self.compiled_output_text.delete("1.0", tk.END)
        self.compiled_output_text.insert(tk.END, text)
        self.compiled_output_text.config(state=tk.DISABLED)
        
        if text.startswith("Error:"):
            self.status_label.config(text="Error occurred.", fg="red")
            messagebox.showerror("Processing Error", text)
        else:
            self.status_label.config(text="Complete.", fg="green")

    def copy_output(self):
        self.master.clipboard_clear()
        self.master.clipboard_append(self.compiled_output_text.get("1.0", tk.END))
        self.status_label.config(text="Copied.")

    def _add_markdown_table_to_doc(self, doc, table_lines):
        if not table_lines: return
        data = []
        for line in table_lines:
            cells = [c.strip() for c in line.strip('|').split('|')]
            data.append(cells)
        if not data: return
        rows = len(data)
        cols = max(len(r) for r in data)
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        # If the second line is a separator like |---|---| treat first row as header
        header_like = False
        if len(table_lines) > 1 and re.match(r'^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+$', table_lines[1].strip()):
            header_like = True
        for i, row in enumerate(data):
            for j, cell_text in enumerate(row):
                if j < len(table.rows[i].cells):
                    cell = table.rows[i].cells[j]
                    cell.text = cell_text
                    if i == 0 and header_like:
                        for run in cell.paragraphs[0].runs:
                            run.font.bold = True

    def export_to_docx(self):
        if not DOCX_AVAILABLE:
            messagebox.showerror("Error", "python-docx not installed")
            return
        
        text = self.compiled_output_text.get("1.0", tk.END).strip()
        if len(text) < 10:
            messagebox.showinfo("Info", "No notes to save.")
            return

        default_filename = f"Notes_{datetime.now().strftime('%Y%m%d')}.docx"
        path = filedialog.asksaveasfilename(
            initialfile=default_filename,
            defaultextension=".docx", 
            filetypes=[("Word Document", "*.docx")]
        )
        
        if not path: return

        doc = Document()
        doc.add_heading("Study Guide", 0)

        lines = text.split('\n')
        in_table = False
        table_buffer = []

        for line in lines:
            line = line.rstrip()
            if line.startswith('|') and '|' in line[1:]:
                # table line
                # ignore separator-only lines except to signal header presence
                if re.match(r'^\|?\s*:-{1,}\s*(\|\s*:-{1,}\s*)+$', line):
                    # treat as separator, don't add
                    in_table = True
                    continue
                table_buffer.append(line)
                in_table = True
                continue
            else:
                if in_table and table_buffer:
                    self._add_markdown_table_to_doc(doc, table_buffer)
                    table_buffer = []
                    in_table = False

            if not line.strip():
                continue
            if line.startswith('## '):
                doc.add_heading(line[3:], level=1)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=2)
            elif line.startswith('* ') or line.startswith('- '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\.\s+', line):
                doc.add_paragraph(re.sub(r'^\d+\.\s+', '', line), style='List Number')
            else:
                doc.add_paragraph(line)

        if in_table and table_buffer:
            self._add_markdown_table_to_doc(doc, table_buffer)

        try:
            doc.save(path)
            messagebox.showinfo("Success", f"Saved to {path}")
        except Exception as e:
            messagebox.showerror("Error", str(e))

if __name__ == "__main__":
    root = tk.Tk()
    app = NoteOrganizerApp(root)
    root.mainloop()
